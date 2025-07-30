# generator.py
from docx import Document
from docx.shared import RGBColor
from docx.oxml.shared import OxmlElement, qn

def generate_doc(
    tables_data,  # Changed from single table to list of tables
    output_file="output_table_doc.docx"
):
    def set_cell_background_color(cell, color):
        """Set background color for a cell"""
        cell_xml_element = cell._tc
        table_cell_properties = cell_xml_element.get_or_add_tcPr()
        shade_obj = OxmlElement('w:shd')
        shade_obj.set(qn('w:fill'), color)
        table_cell_properties.append(shade_obj)
    
    def set_cell_text_color(cell, color):
        """Set text color for a cell"""
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor.from_string(color)
    
    def style_header_cells(cells, bg_color='87CEEB', text_color='FFFFFF'):
        """Apply styling to header cells"""
        for cell in cells:
            set_cell_background_color(cell, bg_color)
            set_cell_text_color(cell, text_color)

    doc = Document()
    
    # Process each table
    for table_index, table_data in enumerate(tables_data, 1):
        table_name = table_data.get('table_name')
        fields = table_data.get('fields')
        create_query = table_data.get('create_query')
        select_query = table_data.get('select_query')
        insert_query = table_data.get('insert_query')
        update_query = table_data.get('update_query')
        delete_query = table_data.get('delete_query')

        # A. Table Description Section
        doc.add_heading(f'{table_index}. Table Description - {table_name}', level=1)
        doc.add_paragraph("Section ini menjelaskan tentang desain table yang terkait dengan project yang akan dilakukan provisioning.")
        
        # Table Name
        doc.add_heading(f"{table_index}.1 Table Name", level=2)
        doc.add_paragraph(f"{table_name}")

        # Structure Table
        doc.add_heading(f"{table_index}.2 Detail Table Field", level=2)
        table = doc.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Field'
        hdr_cells[1].text = 'Data Type'
        hdr_cells[2].text = 'Nullable'
        hdr_cells[3].text = 'Key'
        hdr_cells[4].text = 'Data Pribadi'
        hdr_cells[5].text = 'Description'
        
        # Apply styling to header cells
        style_header_cells(hdr_cells)

        for field in fields:
            row_cells = table.add_row().cells
            row_cells[0].text = field.get("field", "")
            row_cells[1].text = field.get("data_type", "")
            row_cells[2].text = field.get("nullable", "")
            row_cells[3].text = field.get("key", "")
            row_cells[4].text = field.get("data_pribadi", "")
            row_cells[5].text = field.get("description", "")

        # Create Table Query
        doc.add_heading(f"{table_index}.3 Create Table Query", level=2)
        create_table = doc.add_table(rows=2, cols=1)
        create_table.style = 'Table Grid'
        
        # Set header and content
        create_table.rows[0].cells[0].text = "Create Query"
        create_table.rows[1].cells[0].text = create_query
        
        # Apply styling to header cell
        style_header_cells([create_table.rows[0].cells[0]])

        # B. List Query Section
        doc.add_heading(f"{table_index}.4 List Query", level=2)

        def add_query_table(query_index, title, query):
            doc.add_heading(f"{table_index}.4.{query_index} {title}", level=3)
            
            # Query table
            query_table = doc.add_table(rows=2, cols=1)
            query_table.style = 'Table Grid'
            query_table.rows[0].cells[0].text = "Query"
            query_table.rows[1].cells[0].text = query
            style_header_cells([query_table.rows[0].cells[0]])
            
            # Expected Result table
            result_table = doc.add_table(rows=2, cols=1)
            result_table.style = 'Table Grid'
            result_table.rows[0].cells[0].text = "Expected Result"
            result_table.rows[1].cells[0].text = ""
            style_header_cells([result_table.rows[0].cells[0]])
            
            # Cost Before Tuning table
            cost_before_table = doc.add_table(rows=2, cols=1)
            cost_before_table.style = 'Table Grid'
            cost_before_table.rows[0].cells[0].text = "Cost Before Tuning"
            cost_before_table.rows[1].cells[0].text = ""
            style_header_cells([cost_before_table.rows[0].cells[0]])
            
            # Cost After Tuning table
            cost_after_table = doc.add_table(rows=2, cols=1)
            cost_after_table.style = 'Table Grid'
            cost_after_table.rows[0].cells[0].text = "Cost After Tuning"
            cost_after_table.rows[1].cells[0].text = ""
            style_header_cells([cost_after_table.rows[0].cells[0]])

        # Add all query sections with numbering
        add_query_table(1, "Select Query", select_query)
        add_query_table(2, "Insert Query", insert_query)
        add_query_table(3, "Update Query", update_query)
        add_query_table(4, "Delete Query", delete_query)
        
        # Add page break between tables (except for the last one)
        if table_index < len(tables_data):
            doc.add_page_break()

    doc.save(output_file)
    print(f"✅ Document saved as: {output_file}")
