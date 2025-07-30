from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# Function to set the style for table headers
def set_table_header_style(row):
    """
    Set background color #cfe1f3, white text, bold, Tahoma, size 11 for all cells in the given row.
    """
    for cell in row.cells:
        # Set background color
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), 'cfe1f3')
        tcPr.append(shd)
        # Set text style
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.name = 'Tahoma'
                run.font.size = Pt(11)
                r = run._element
                r.rPr.rFonts.set(qn('w:eastAsia'), 'Tahoma')

# Function to style heading level 1 and 2
def set_heading_style(paragraph):
    from docx.shared import RGBColor
    for run in paragraph.runs:
        run.bold = True
        run.font.name = 'Tahoma'
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0, 0, 0)
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), 'Tahoma')

# Function to style normal content (not table header, heading, or special table row)
def set_content_style(paragraph):
    from docx.shared import RGBColor
    for run in paragraph.runs:
        run.font.name = 'Tahoma'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0, 0, 0)
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), 'Tahoma')

def set_table_content_style(table, header_rows=None, special_rows=None):
    """
    Apply set_content_style to all cells in the table except header_rows and special_rows.
    header_rows and special_rows should be lists of row indices.
    """
    if header_rows is None:
        header_rows = [0]
    if special_rows is None:
        special_rows = []
    for i, row in enumerate(table.rows):
        if i in header_rows or i in special_rows:
            continue
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                set_content_style(paragraph)


# Service Table
heading = doc.add_heading('2.1 Generate Access Token', level=1)
set_heading_style(heading)
table = doc.add_table(rows=2, cols=2)
table.style = 'Table Grid'
table.cell(0, 0).text = 'Nama Service'
table.cell(0, 1).text = 'Fungsi'
table.cell(1, 0).text = 'Generate Access Token'
table.cell(1, 1).text = (
    'API yang digunakan untuk generate clientSecret (Access token menggunakan RSA Key Pair 515 bit dengan request clientId yang terdaftar dan kemudian dilakukan encrypt menggunakan base64 dengan kombinasi privateKey dan juga secretKey dari random string, clientSecret (Access token) akan tersimpan pada repository redis dan memiliki masa berlaku selama 3 jam. Anda dapat membuat access token berulang kali tanpa membuat access token sebelumnya menjadi tidak berlaku, selama belum kedaluwarsa (belum lebih dari 3 jam)'
)
# Set header style for service_table
set_table_header_style(table.rows[0])
set_table_content_style(table)

# HTTP Request
heading = doc.add_heading('1. HTTP Request', level=2)
set_heading_style(heading)
table = doc.add_table(rows=2, cols=2)
table.style = 'Table Grid'
table.cell(0, 0).text = 'Method'
table.cell(0, 1).text = 'URL/Endpoint'
table.cell(1, 0).text = 'POST'
table.cell(1, 1).text = '<context-path>/api/v1/switcher/data-processing/generate/access-token/merchant'
# Set header style for HTTP Request table
set_table_header_style(table.rows[0])
set_table_content_style(table)

# Authorization
heading = doc.add_heading('2. Authorization', level=2)
set_heading_style(heading)
desc = doc.add_paragraph('Token merupakan hasil encrypt dari private key dan secret key(random string) yang menggunakan format Base64. Masa berlaku token ini adalah 1 jam. Anda dapat membuat token berulang kali tanpa membuat token sebelumnya menjadi tidak berlaku, selama belum kadaluarsa (belum lebih dari 1 jam).')
set_content_style(desc)
table = doc.add_table(rows=2, cols=2)
table.style = 'Table Grid'
table.autofit = False
table.cell(0, 0).text = 'Authorization'
table.cell(0, 1).text = 'Bearer'
table.cell(1, 0).text = 'POST'
table.cell(1, 1).text = 'eyJhbGciOiJIUzI1NiJ9.eyJzdWIiOiJUVkpEVGpBd01nPT0iLCJpYXQiOjE3MjIzMDI5OTQsImV4cCI6MTcyMjMzODk5NH0.oloUIbs8NbhqgG4YYHFg4-B53xEH5CHZ9mUjtbUkSGs'
# Set header style for Authorization table
set_table_header_style(table.rows[0])
set_table_content_style(table)

# Headers
heading = doc.add_heading('3. Headers', level=2)
set_heading_style(heading)
table = doc.add_table(rows=2, cols=2)
table.style = 'Table Grid'
table.autofit = False
table.cell(0, 0).text = 'Header Key'
table.cell(0, 1).text = 'Header Value'
table.cell(1, 0).text = 'Authorization'
table.cell(1, 1).text = 'Bearer eyJhbGciOiJIUzI1NiJ9.eyJzdWIiOiJUVkpEVGpBd01nPT0iLCJpYXQiOjE3MjIzMDI5OTQsImV4cCI6MTcyMjMzODk5NH0.oloUIbs8NbhqgG4YYHFg4-B53xEH5CHZ9mUjtbUkSGs'
# Set header style for Headers table
set_table_header_style(table.rows[0])
set_table_content_style(table)

# Parameters
heading = doc.add_heading('4. Parameters', level=2)
set_heading_style(heading)
table = doc.add_table(rows=2, cols=5)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Parameters'
hdr_cells[1].text = 'Data Type'
hdr_cells[2].text = 'Length'
hdr_cells[3].text = 'Mandatory'
hdr_cells[4].text = 'Description'
row_cells = table.rows[1].cells
row_cells[0].text = 'clientId'
row_cells[1].text = 'Text'
row_cells[2].text = '50'
row_cells[3].text = 'Mandatory'
row_cells[4].text = 'kode client id untuk esb'
# Set header style for Parameters table
set_table_header_style(table.rows[0])
set_table_content_style(table)

# Result
heading = doc.add_heading('5. Result', level=2)
set_heading_style(heading)
table = doc.add_table(rows=5, cols=2)
table.style = 'Table Grid'
table.cell(0, 0).text = 'Parameters'
table.cell(0, 1).text = 'Description'
table.cell(1, 0).text = 'Status'
table.cell(1, 1).text = 'Representasi dari hasil request API (true/false)'
table.cell(2, 0).text = 'Message'
table.cell(2, 1).text = 'Deskripsi atau keterangan sesuai dengan response code yang dikirim server'
table.cell(3, 0).text = 'Data'
table.cell(3, 1).text = 'Informasi lengkap yang dikirim server terhadap request API'
table.cell(4, 0).text = 'Paging'
table.cell(4, 1).text = 'Informasi paging jika ada, berisi total data, total halaman, halaman saat ini, dan jumlah data per halaman'
# Set header style for Result table
set_table_header_style(table.rows[0])
set_table_content_style(table)

# Example JSON
heading = doc.add_heading('6. Example Parameter With JSON', level=2)
set_heading_style(heading)
table = doc.add_table(rows=6, cols=1)
table.style = 'Table Grid'
table.cell(0, 0).text = 'Example JSON Request Body'
table.cell(1, 0).text = (
    '{\n'
    '  "clientId": "MRCN002-switcher-app"\n'
    '}'
)
table.cell(2, 0).text = 'JSON Response Body Success'
table.cell(3, 0).text = (
    '{\n'
    '  "success": true,\n'
    '  "responseCode": "SCS01200",\n'
    '  "responseMessage": "success",\n'
    '  "data": {\n'
    '    "accessToken": "Bb6DbNns3IC69C8DCn7TD9iIYuB0pOslMkY+h/DkQ1b/0fy3ID6WpNOc67TjchThFfFuAzkg+kJOAXtjgP3mzA==",\n'
    '    "tokenType": "bearer",\n'
    '    "expiresIn": "10800",\n'
    '    "expiresAt": "2024-09-13T01:28:12+07:00"\n'
    '  },\n'
    '  "version": "1.0"\n'
    '}'
)
table.cell(4, 0).text = 'JSON Response Body Error'
table.cell(5, 0).text = (
    '{\n'
    '  "success": false,\n'
    '  "responseCode": "ERR01404",\n'
    '  "responseMessage": "client id not registered on the payment switcher app",\n'
    '  "data": null,\n'
    '  "version": "1.0"\n'
    '}'
)
# Set header style for Example JSON table (rows 0, 2, 4)
for i in [0, 2, 4]:
    set_table_header_style(table.rows[i])
    
set_table_content_style(table, header_rows=[0, 2, 4], special_rows=[1, 3, 5])

doc.save('Generate_Access_Token_API.docx')
print("Word document generated: Generate_Access_Token_API.docx")