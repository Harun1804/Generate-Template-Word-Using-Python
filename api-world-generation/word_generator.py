from docx import Document
from word_styles import WordStyles

class WordTemplateGenerator:
    def __init__(self, data):
        self.data = data
        self.doc = Document()

    def add_api_section(self, api):
        main_section = 2  # or whatever your main section number is
        for idx, api in enumerate(self.data.apis):
            heading = self.doc.add_heading(f"{main_section}.{idx+1} {api['service_title']}", level=1)
            WordStyles.set_heading_style(heading)
        
        # Service Table
        WordStyles.set_heading_style(heading)
        table = self.doc.add_table(rows=2, cols=2)
        table.style = 'Table Grid'
        for col, val in enumerate(api["service_table"]["header"]):
            table.cell(0, col).text = val
        for col, val in enumerate(api["service_table"]["rows"][0]):
            table.cell(1, col).text = val
        WordStyles.set_table_header_style(table.rows[0])
        WordStyles.set_table_content_style(table)

        # Section numbering
        section_num = 1

        # HTTP Request
        heading = self.doc.add_heading(f"{section_num}. HTTP Request", level=2)
        WordStyles.set_heading_style(heading)
        section_num += 1
        table = self.doc.add_table(rows=2, cols=2)
        table.style = 'Table Grid'
        for col, val in enumerate(api["http_request"]["header"]):
            table.cell(0, col).text = val
        for col, val in enumerate(api["http_request"]["rows"][0]):
            table.cell(1, col).text = val
        WordStyles.set_table_header_style(table.rows[0])
        WordStyles.set_table_content_style(table)

        # Optional: Authorization
        if "authorization_desc" in api and "authorization_table" in api:
            heading = self.doc.add_heading(f"{section_num}. Authorization", level=2)
            WordStyles.set_heading_style(heading)
            para = self.doc.add_paragraph(api["authorization_desc"])
            WordStyles.set_content_style(para)
            table = self.doc.add_table(rows=2, cols=2)
            table.style = 'Table Grid'
            for col, val in enumerate(api["authorization_table"]["header"]):
                table.cell(0, col).text = val
            for col, val in enumerate(api["authorization_table"]["rows"][0]):
                table.cell(1, col).text = val
            WordStyles.set_table_header_style(table.rows[0])
            WordStyles.set_table_content_style(table)
            section_num += 1

        # Optional: Headers
        if "headers_table" in api:
            heading = self.doc.add_heading(f"{section_num}. Headers", level=2)
            WordStyles.set_heading_style(heading)
            table = self.doc.add_table(rows=2, cols=2)
            table.style = 'Table Grid'
            for col, val in enumerate(api["headers_table"]["header"]):
                table.cell(0, col).text = val
            for col, val in enumerate(api["headers_table"]["rows"][0]):
                table.cell(1, col).text = val
            WordStyles.set_table_header_style(table.rows[0])
            WordStyles.set_table_content_style(table)
            section_num += 1

        # Parameters
        heading = self.doc.add_heading(f"{section_num}. Parameters", level=2)
        WordStyles.set_heading_style(heading)
        table = self.doc.add_table(rows=2, cols=5)
        table.style = 'Table Grid'
        for col, val in enumerate(api["parameters_table"]["header"]):
            table.cell(0, col).text = val
        for col, val in enumerate(api["parameters_table"]["rows"][0]):
            table.cell(1, col).text = val
        WordStyles.set_table_header_style(table.rows[0])
        WordStyles.set_table_content_style(table)
        section_num += 1

        # Result
        heading = self.doc.add_heading(f"{section_num}. Result", level=2)
        WordStyles.set_heading_style(heading)
        table = self.doc.add_table(rows=len(api["result_table"]["rows"])+1, cols=2)
        table.style = 'Table Grid'
        for j, val in enumerate(api["result_table"]["header"]):
            table.cell(0, j).text = val
        for i, row in enumerate(api["result_table"]["rows"]):
            for j, val in enumerate(row):
                table.cell(i+1, j).text = val
        WordStyles.set_table_header_style(table.rows[0])
        WordStyles.set_table_content_style(table)
        section_num += 1

        # Example JSON
        heading = self.doc.add_heading(f"{section_num}. Example Parameter With JSON", level=2)
        WordStyles.set_heading_style(heading)
        rows = api["example_json_table"]["rows"]
        table = self.doc.add_table(rows=len(rows), cols=1)
        table.style = 'Table Grid'
        for i, row in enumerate(rows):
            table.cell(i, 0).text = row[0]
        for i in api["example_json_table"]["header_rows"]:
            WordStyles.set_table_header_style(table.rows[i])
        WordStyles.set_table_content_style(
            table,
            header_rows=api["example_json_table"]["header_rows"],
            special_rows=api["example_json_table"]["special_rows"]
        )

    def save(self, filename):
        self.doc.save(filename)

    def generate(self):
        for api in self.data.apis:
            self.add_api_section(api)